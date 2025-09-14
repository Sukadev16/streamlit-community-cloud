"""
ai_data_storyteller_step_by_step.py

Step-by-step, well-documented Streamlit dashboard for the assessment:
- CSV upload and validation
- Automated EDA (summary statistics, missing values, value counts, correlations)
- Templated natural-language insights (with optional local LLM hook)
- At least 3 visualizations: bar chart (categorical), line/trend chart (numeric/time), correlation heatmap/hist
- Generate a concise 1-2 page executive summary as PDF or Word

How to use:
1. Create a Python environment and install dependencies (example):
   pip install streamlit pandas numpy matplotlib seaborn plotly python-docx fpdf pillow scikit-learn
   # Optional for LLM support:
   pip install transformers

2. Run the app:
   streamlit run ai_data_storyteller_step_by_step.py

Notes:
- This file is intentionally verbose and educational: each function is documented and each Streamlit UI block is labelled as a step.
- If you want a Jupyter notebook, you can convert using jupytext or paste cells into a new notebook. I can also produce the .ipynb for you on request.
"""

# --------------------------- Imports ---------------------------
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import os
import io
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from datetime import datetime

# Optional LLM (disabled by default). To enable set USE_LLM = True and provide a compatible local model name.
USE_LLM = False
LLM_MODEL_NAME = "gpt4all"  # placeholder
if USE_LLM:
    try:
        from transformers import pipeline
        llm_pipe = pipeline("text-generation", model=LLM_MODEL_NAME)
    except Exception as e:
        # If LLM initialization fails, fall back to templated insights and warn in UI
        llm_pipe = None
        USE_LLM = False

# --------------------------- Helpers: File & Validation ---------------------------

def read_csv_file(uploaded_file):
    """Read uploaded CSV into a DataFrame with safe defaults.

    Returns (df, error_message) where error_message is None when successful.
    """
    try:
        # try pandas default read
        df = pd.read_csv(uploaded_file)
        return df, None
    except Exception as e:
        # try with more forgiving options
        try:
            df = pd.read_csv(uploaded_file, encoding='utf-8', engine='python')
            return df, None
        except Exception as e2:
            return None, f"Error reading CSV: {e}; fallback error: {e2}"


def validate_dataframe(df: pd.DataFrame):
    """Perform a few pragmatic validation checks and return list of issues (empty if valid).

    Checks:
    - Not empty
    - Has >= 2 rows
    - Has column names (no empty column name)
    - Reasonable memory size
    """
    problems = []
    if df is None:
        problems.append("No dataframe provided")
        return problems
    if df.empty:
        problems.append("Dataset is empty")
    if df.shape[0] < 2:
        problems.append("Dataset has fewer than 2 rows")
    if df.shape[1] < 1:
        problems.append("Dataset has no columns")
    # Column name checks
    for c in df.columns:
        if str(c).strip() == "":
            problems.append("One or more columns have empty names")
            break
    # memory check (avoid extremely large uploads in this assessment context)
    try:
        mem_bytes = df.memory_usage(deep=True).sum()
        if mem_bytes > 500e6:  # 500 MB limit for interactive demo
            problems.append("Dataset is very large (>500MB) — consider sampling or uploading a smaller file")
    except Exception:
        pass
    return problems

# --------------------------- EDA Functions ---------------------------

def summary_statistics(df: pd.DataFrame):
    """Return a descriptive summary of numeric columns and an extended summary for all columns."""
    numeric_summary = df.describe().T
    full_summary = df.describe(include='all').T
    return numeric_summary, full_summary


def missing_values(df: pd.DataFrame):
    """Return missing values counts and percent by column."""
    miss_count = df.isnull().sum()
    miss_pct = 100 * miss_count / len(df)
    miss_df = pd.DataFrame({'missing_count': miss_count, 'missing_pct': miss_pct})
    miss_df = miss_df.sort_values('missing_pct', ascending=False)
    return miss_df


def value_counts_for_categoricals(df: pd.DataFrame, top_n=10):
    cat_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
    result = {}
    for c in cat_cols:
        result[c] = df[c].value_counts(dropna=False).head(top_n)
    return result


def compute_correlations(df: pd.DataFrame):
    num_cols = df.select_dtypes(include=['number']).columns.tolist()
    if not num_cols:
        return None
    corr = df[num_cols].corr()
    return corr

# --------------------------- Insigh Generation (templated + optional LLM) ---------------------------

def generate_basic_insights(df: pd.DataFrame, n_top=5):
    """Create templated plain-English insights suitable for executive summary.

    This function is deterministic and does not require an LLM. It extracts the most useful signals.
    """
    insights = []

    insights.append(f"The dataset contains {df.shape[0]:,} rows and {df.shape[1]} columns.")

    # missing info
    miss_total = int(df.isnull().sum().sum())
    miss_pct = 100 * miss_total / (df.shape[0] * max(df.shape[1], 1))
    insights.append(f"There are {miss_total} missing values ({miss_pct:.2f}% of all cells).")

    num_cols = df.select_dtypes(include=['number']).columns.tolist()
    cat_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()

    if num_cols:
        # Show top columns by variance (may indicate important features)
        variances = df[num_cols].var().sort_values(ascending=False)
        top_var = variances.head(n_top)
        insights.append(f"Top numeric columns by variance: {', '.join([f'{c} (var={v:.2f})' for c,v in top_var.items()])}.")

        # strong correlation pair
        corr = df[num_cols].corr().abs()
        if corr.shape[0] > 1:
            corr_unstack = corr.where(~np.eye(corr.shape[0],dtype=bool)).stack()
            if not corr_unstack.empty:
                top_idx = corr_unstack.idxmax()
                top_val = corr_unstack.max()
                insights.append(f"Strong numeric correlation detected between '{top_idx[0]}' and '{top_idx[1]}' (|r| ≈ {top_val:.2f}).")
    else:
        insights.append("No numeric columns found for variance/correlation analysis.")

    if cat_cols:
        sample = []
        for c in cat_cols[:n_top]:
            vc = df[c].value_counts(dropna=False).head(3)
            sample.append(f"{c}: {', '.join([f'{idx}({cnt})' for idx,cnt in vc.items()])}")
        insights.append("Top categorical distributions — " + "; ".join(sample))
    else:
        insights.append("No categorical columns detected.")

    # Detect obvious outliers metric-wise (IQR rule on numeric columns)
    outlier_info = []
    for c in num_cols[:n_top]:
        series = df[c].dropna()
        if series.empty:
            continue
        q1 = series.quantile(0.25)
        q3 = series.quantile(0.75)
        iqr = q3 - q1
        if iqr == 0:
            continue
        lower = q1 - 1.5 * iqr
        upper = q3 + 1.5 * iqr
        n_out = ((series < lower) | (series > upper)).sum()
        if n_out > 0:
            outlier_info.append(f"{c} has {n_out} outlier(s)")
    if outlier_info:
        insights.append("Outlier note: " + ", ".join(outlier_info))

    return insights


def llm_generate_insights(prompt: str, max_length=256):
    """Generate insights using an LLM pipeline (if available).

    Returns generated text or None if LLM not available.
    """
    if not USE_LLM or 'llm_pipe' not in globals() or llm_pipe is None:
        return None
    try:
        out = llm_pipe(prompt, max_length=max_length, do_sample=False)
        # Huggingface pipeline returns list of dicts: [{'generated_text': '...'}]
        return out[0].get('generated_text', None)
    except Exception as e:
        return None

# --------------------------- Visualization Helpers ---------------------------

def save_fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf


def plot_bar_top_categories(df: pd.DataFrame, col: str, top_n=10):
    fig, ax = plt.subplots(figsize=(6,4))
    vc = df[col].value_counts(dropna=False).head(top_n)
    vc.plot(kind='bar', ax=ax)
    ax.set_title(f"Top {top_n} values: {col}")
    ax.set_xlabel(col)
    ax.set_ylabel('count')
    plt.tight_layout()
    return fig


def plot_line_trend(df: pd.DataFrame, x_col: str, y_col: str):
    fig, ax = plt.subplots(figsize=(8,4))
    ax.plot(df[x_col], df[y_col], marker='o', linestyle='-')
    ax.set_title(f"{y_col} over {x_col}")
    ax.set_xlabel(x_col)
    ax.set_ylabel(y_col)
    plt.xticks(rotation=30)
    plt.tight_layout()
    return fig


def plot_correlation_heatmap(df: pd.DataFrame):
    num = df.select_dtypes(include=['number'])
    fig, ax = plt.subplots(figsize=(7,6))
    sns.heatmap(num.corr(), annot=True, fmt='.2f', cmap='coolwarm', ax=ax)
    ax.set_title('Correlation heatmap')
    plt.tight_layout()
    return fig


def plot_histograms(df: pd.DataFrame, columns=None, bins=20):
    if columns is None:
        columns = df.select_dtypes(include=['number']).columns.tolist()
    n = len(columns)
    if n == 0:
        return None
    cols = min(n, 4)
    rows = (n + cols - 1) // cols
    fig, axes = plt.subplots(rows, cols, figsize=(cols*4, rows*3))
    axes = np.array(axes).reshape(-1)
    for ax, col in zip(axes, columns):
        ax.hist(df[col].dropna(), bins=bins)
        ax.set_title(col)
    # hide unused axes
    for ax in axes[n:]:
        ax.axis('off')
    plt.tight_layout()
    return fig

# --------------------------- Report Export (PDF and DOCX) ---------------------------

def make_pdf_report(title: str, insights: list, chart_bytes_list: list, out_path: str):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, title, ln=True, align='C')
    pdf.ln(5)
    pdf.set_font('Arial', size=12)
    pdf.multi_cell(0, 6, 'Executive Summary:')
    pdf.ln(2)
    for i, ins in enumerate(insights, 1):
        pdf.multi_cell(0, 6, f"{i}. {ins}")
    # add charts each on new page
    for i, buf in enumerate(chart_bytes_list):
        try:
            tmp_path = f"_tmp_chart_{i}.png"
            with open(tmp_path, 'wb') as f:
                f.write(buf.getbuffer())
            pdf.add_page()
            pdf.image(tmp_path, x=15, w=180)
            os.remove(tmp_path)
        except Exception as e:
            print('Error adding chart to PDF', e)
    pdf.output(out_path)


def make_docx_report(title: str, insights: list, chart_bytes_list: list, out_path: str):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_heading('Executive Summary', level=2)
    for ins in insights:
        doc.add_paragraph(ins)
    doc.add_heading('Visualizations', level=2)
    for i, buf in enumerate(chart_bytes_list):
        tmp_path = f"_tmp_chart_doc_{i}.png"
        with open(tmp_path, 'wb') as f:
            f.write(buf.getbuffer())
        try:
            doc.add_picture(tmp_path, width=Inches(6))
        except Exception as e:
            print('Could not add picture to docx', e)
        os.remove(tmp_path)
    doc.save(out_path)

# --------------------------- Streamlit UI: Step-by-step ---------------------------

def main():
    st.set_page_config(page_title='AI Data Storyteller - Step by Step', layout='wide')

    st.title('📊 AI Data Storyteller — Step-by-step')
    st.markdown(
        "This interactive dashboard shows step-by-step EDA, auto insights, visualizations, and report export."
    )

    # ---------------- Step 1: Data input ----------------
    st.header('Step 1 — Dataset upload & validation')
    with st.expander('Upload instructions (click to view)'):
        st.write(
            'Upload a CSV file. The app validates basic sanity checks (not empty, -- at least 2 rows, named columns, reasonable size).'
        )

    col1, col2 = st.columns([2,1])
    with col1:
        uploaded_file = st.file_uploader('Upload CSV file', type=['csv'])
        if uploaded_file is not None:
            df, err = read_csv_file(uploaded_file)
            if err:
                st.error(err)
                return
            st.success('CSV loaded successfully')
            st.write('First 10 rows:')
            st.dataframe(df.head(10))
        else:
            st.info('No file uploaded — you may use a sample dataset from the sidebar')
            df = None

    with col2:
        st.write('Sample data')
        sample = st.selectbox('Load sample', options=['None', 'Iris (sklearn)', 'Tips (seaborn)'])
        if sample != 'None' and df is None:
            if sample == 'Iris (sklearn)':
                from sklearn import datasets
                iris = datasets.load_iris()
                df = pd.DataFrame(iris.data, columns=iris.feature_names)
                df['target'] = iris.target
            elif sample == 'Tips (seaborn)':
                df = sns.load_dataset('tips')
            st.success(f'Loaded sample: {sample}')

    # If upload or sample provided, validate and continue
    if df is not None:
        problems = validate_dataframe(df)
        if problems:
            st.error('Validation problems found:')
            for p in problems:
                st.write('- ' + p)
            st.stop()
        else:
            st.success('Dataset validation passed')

        # ---------------- Step 2: Automated EDA ----------------
        st.header('Step 2 — Automated EDA')
        st.subheader('Summary statistics (numeric)')
        num_summary, full_summary = summary_statistics(df)
        st.dataframe(num_summary)

        st.subheader('Missing values (by column)')
        miss_df = missing_values(df)
        st.dataframe(miss_df[miss_df['missing_count']>0])

        st.subheader('Value counts (categorical)')
        vc = value_counts_for_categoricals(df)
        if vc:
            for c, series in vc.items():
                st.write(f'Column: {c}')
                st.write(series)
        else:
            st.info('No categorical columns detected')

        st.subheader('Correlation (numeric)')
        corr = compute_correlations(df)
        if corr is not None:
            st.write(corr)
            fig_corr = plot_correlation_heatmap(df)
            st.pyplot(fig_corr)
        else:
            st.info('Not enough numeric columns for correlation analysis')

        # ---------------- Step 3: Generate insights ----------------
        st.header('Step 3 — Generate plain-English insights')
        use_llm = st.checkbox('Use LLM (local) to generate insights (optional)', value=False)
        insights = None
        if use_llm and USE_LLM:
            # build the prompt (short)
            prompt = f"You are a data analyst. Provide top 6 concise business insights for the dataset with {df.shape[0]} rows and {df.shape[1]} columns. Columns: {', '.join(df.columns[:40])}."
            text = llm_generate_insights(prompt)
            if text:
                st.text_area('LLM insights', value=text, height=200)
                insights = [line.strip() for line in text.split('\n') if line.strip()]
            else:
                st.warning('LLM generation failed — falling back to templated insights')
        if insights is None:
            insights = generate_basic_insights(df, n_top=5)
            for i, ins in enumerate(insights, 1):
                st.write(f'{i}. {ins}')

        # ---------------- Step 4: Auto Visualizations ----------------
        st.header('Step 4 — Auto visualizations (at least 3)')
        st.write('The app will attempt to generate: 1) bar chart (categorical), 2) line/trend chart (time or index), 3) histograms/heatmap.')

        chart_buffers = []

        # 1) Bar chart for top categorical (if any)
        cat_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        if cat_cols:
            chosen_cat = cat_cols[0]
            st.subheader(f'Bar chart — top values for {chosen_cat}')
            fig_bar = plot_bar_top_categories(df, chosen_cat)
            st.pyplot(fig_bar)
            chart_buffers.append(save_fig_to_bytes(fig_bar))
        else:
            st.info('No categorical columns found for bar chart')

        # 2) Line chart: prefer datetime + numeric, else two numeric columns or numeric over index
        date_cols = df.select_dtypes(include=['datetime', 'datetime64']).columns.tolist()
        # detect date-like strings and convert
        if not date_cols:
            for col in df.columns:
                if df[col].dtype == 'object':
                    try:
                        parsed = pd.to_datetime(df[col], errors='coerce')
                        if parsed.notnull().sum() > 0.5*len(parsed):
                            df[col] = parsed
                            date_cols.append(col)
                            break
                    except Exception:
                        continue
        if date_cols and df.select_dtypes(include=['number']).shape[1] > 0:
            dcol = date_cols[0]
            ycol = df.select_dtypes(include=['number']).columns[0]
            st.subheader(f'Line chart — {ycol} over {dcol}')
            df_sorted = df.sort_values(dcol)
            fig_line = plot_line_trend(df_sorted, dcol, ycol)
            st.pyplot(fig_line)
            chart_buffers.append(save_fig_to_bytes(fig_line))
        elif df.select_dtypes(include=['number']).shape[1] >= 2:
            cols = df.select_dtypes(include=['number']).columns[:2]
            st.subheader(f'Line plot (index) — {cols[0]} and {cols[1]} over index')
            fig_line = plt.figure(figsize=(8,4))
            plt.plot(df[cols[0]].values, label=str(cols[0]))
            plt.plot(df[cols[1]].values, label=str(cols[1]))
            plt.legend()
            plt.title('Numeric trends over index')
            st.pyplot(fig_line)
            # convert current figure
            fig_line_buf = io.BytesIO()
            plt.savefig(fig_line_buf, format='png', bbox_inches='tight')
            fig_line_buf.seek(0)
            chart_buffers.append(fig_line_buf)
            plt.close()
        else:
            st.info('Not enough numeric/date data for a line chart')

        # 3) Histograms for numeric columns
        num_cols = df.select_dtypes(include=['number']).columns.tolist()
        if num_cols:
            st.subheader('Numeric distributions — histograms')
            fig_hist = plot_histograms(df, columns=num_cols[:4])
            if fig_hist is not None:
                st.pyplot(fig_hist)
                chart_buffers.append(save_fig_to_bytes(fig_hist))
        else:
            st.info('No numeric columns for histograms')

        # At this point we should have collected at least 2–3 charts in chart_buffers
        st.write(f'Collected {len(chart_buffers)} chart(s) for report')

        # ---------------- Step 5: Report generation ----------------
        st.header('Step 5 — Executive summary report')
        report_title = st.text_input('Report title', value=f'Executive Summary - {datetime.now().date()}')
        include_charts = st.checkbox('Include charts in report', value=True)

        # Button to generate Word
        col_docx = st.columns(1)[0] # Use a single column for the remaining button

        with col_docx:
            if st.button('Generate Word (.docx)'):
                tmp_docx = f"executive_summary_{int(datetime.now().timestamp())}.docx"
                chart_list = chart_buffers if include_charts else []
                make_docx_report(report_title, insights, chart_list, tmp_docx)
                with open(tmp_docx, 'rb') as f:
                    data = f.read()
                st.success('Word document generated')
                st.download_button('Download .docx', data=data, file_name='executive_summary.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                os.remove(tmp_docx)

    else:
        st.info('Upload a CSV or select a sample to begin the analysis.')

# --------------------------- Entrypoint ---------------------------

if __name__ == '__main__':
    main()